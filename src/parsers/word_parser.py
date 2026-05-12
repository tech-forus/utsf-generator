"""
Word Document Parser
====================
Extracts text and tables from .docx (Word) files using python-docx.

Note: .doc (old Word 97-2003) is NOT supported — open in Word and
      Save As .docx first.
"""

import os
from typing import Dict, Any, List, Optional
from parsers.base_parser import BaseParser


class WordParser(BaseParser):
    SUPPORTED_EXTENSIONS = [".docx"]

    def parse(self, file_path: str, doc_context=None) -> Dict[str, Any]:
        try:
            from docx import Document
        except ImportError:
            print("[WordParser] python-docx not installed — pip install python-docx")
            return {"text": "", "tables": [], "data": {}}

        doc = Document(file_path)

        # ── Extract paragraphs ─────────────────────────────────────────────────
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text_body = "\n".join(paragraphs)

        # ── Extract tables ─────────────────────────────────────────────────────
        raw_tables = []
        table_text = ""
        sheets: Dict[str, List[List]] = {}

        for i, tbl in enumerate(doc.tables):
            rows = []
            for row in tbl.rows:
                cells = [cell.text.strip() for cell in row.cells]
                # Merge duplicate adjacent cells (Word table merging artefact)
                deduped = []
                for c in cells:
                    if not deduped or c != deduped[-1]:
                        deduped.append(c)
                if any(deduped):
                    rows.append(deduped)

            if rows:
                name = f"Table {i + 1}"
                raw_tables.append({"name": name, "rows": rows})
                sheets[name] = rows
                table_text += f"\n=== {name} ===\n"
                for row in rows[:100]:
                    table_text += "\t".join(row) + "\n"

        full_text = text_body
        if table_text:
            full_text += "\n\n--- Tables ---\n" + table_text

        # ── Auto-detect structured data from tables ────────────────────────────
        data = {}
        if sheets:
            try:
                from parsers.excel_parser import ExcelParser
                data = ExcelParser()._auto_detect(sheets)
            except Exception:
                pass

        # ── Charge extraction from paragraph text ─────────────────────────────
        # Word docs often have charges as numbered/bulleted lists in paragraph
        # text, not in tables. Run PDF-style free-text extraction on body text.
        if text_body:
            try:
                from parsers.pdf_parser import PDFParser
                _pp = PDFParser.__new__(PDFParser)
                # Section-filter first to avoid T&C ghost values
                try:
                    from knowledge.section_segmenter import SectionSegmenter
                    _seg = SectionSegmenter()
                    _sections = _seg.segment_text(text_body)
                    _sm = _seg.get_sections_map(_sections)
                    _charge_text = ""
                    for _cat in ("CHARGES", "MIXED", "UNKNOWN", "ZONE_MATRIX"):
                        for _sec in _sm.get(_cat, []):
                            _charge_text += "\n" + _sec.text
                    if not _charge_text.strip():
                        _charge_text = text_body
                except Exception:
                    _charge_text = text_body

                text_charges = _pp._extract_charges_from_text(_charge_text)
                if text_charges:
                    data.setdefault("charges", {})
                    for k, v in text_charges.items():
                        data["charges"].setdefault(k, v)

                # Company info from paragraph text (GST, PAN, phone, email)
                cd_text = _pp._extract_company_from_text(text_body)
                if cd_text:
                    data.setdefault("company_details", {})
                    for k, v in cd_text.items():
                        if k not in data["company_details"]:
                            data["company_details"][k] = v
            except Exception as _wp_err:
                print(f"[WordParser] Text extraction error: {_wp_err}")

        return {
            "text": full_text,
            "tables": raw_tables,
            "data": data,
        }
